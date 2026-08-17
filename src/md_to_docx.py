"""TEKNIK_RAPOR.md -> TEKNIK_RAPOR.docx donusumu.

Neden var: .docx onceden ELLE uretiliyordu ve rapor her yeniden uretildiginde
bayatliyordu (teslim/README.md'de uyari olarak duruyordu). Bu betik donusumu
yeniden uretilebilir kilar.

Donusum kurallari, mevcut .docx'ten cikarilmistir ve BILEREK birebir korunur:
    #      -> Title            ###   -> Heading 3        - x  -> List Bullet
    ##     -> Heading 1        > x   -> Normal           1. x -> List Number
    | .. | -> Table Grid (baslik satiri kalin, 9 punto)
    **x**  -> kalin run        `x`   -> ters tirnak atilir, duz metin
Kaynak satirlar TEK TEK paragrafa cevrilir; belgenin mevcut gorunumu boyle
oldugu icin satir birlestirme YAPILMAZ.

Kullanim:
    python -m src.md_to_docx                       # teslim/TEKNIK_RAPOR.md -> .docx
    python -m src.md_to_docx --md X.md --out Y.docx
"""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

from src import config

CELL_PT = Pt(9)


class BoldState:
    """Kalin isaretinin SATIR SINIRINI asmasini yonetir.

    Markdown kaynagi sabit genislikte sarildigi icin `**...**` bir sonraki
    satira tasabiliyor. Satir satir donusturuldugunde bu, metinde ciplak `**`
    birakiyordu. Durumu satirlar arasi tasiyarak hem yildizlar temizleniyor hem
    de kalinlik dogru araliga uygulaniyor.
    """

    def __init__(self):
        self.bold = False

    def apply(self, par, text: str) -> None:
        text = text.replace("`", "")
        for i, part in enumerate(text.split("**")):
            if i:
                self.bold = not self.bold
            if part:
                par.add_run(part).bold = self.bold or None


def split_row(line: str) -> list:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    state = BoldState()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- tablo blogu ---
        if stripped.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            header = split_row(stripped)
            i += 2
            body = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(body), cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, cells in enumerate([header] + body):
                for k in range(len(header)):
                    cell = table.cell(j, k)
                    cell.text = ""
                    par = cell.paragraphs[0]
                    cell_state = BoldState()
                    cell_state.apply(par, cells[k] if k < len(cells) else "")
                    for r in par.runs:
                        r.font.size = CELL_PT
                        if j == 0:
                            r.bold = True
            continue

        if not stripped or stripped == "---":
            # Ne bos satir ne yatay cizgi paragraf uretir; mevcut .docx boyle
            # uretilmisti ve bos paragraf eklemek belgeyi gereksiz sisiriyor.
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            body = stripped[level:].strip()
            style = {1: "Title", 2: "Heading 1"}.get(level, "Heading 3")
            state.apply(doc.add_paragraph(style=style), body)
        elif stripped.startswith("> "):
            state.apply(doc.add_paragraph(style="Normal"), stripped[2:])
        elif stripped.startswith("- "):
            state.apply(doc.add_paragraph(style="List Bullet"), stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            state.apply(doc.add_paragraph(style="List Number"),
                        re.sub(r"^\d+\.\s+", "", stripped))
        else:
            state.apply(doc.add_paragraph(style="Normal"), stripped)
        i += 1

    doc.save(out_path)
    print(f"yazildi: {out_path}  ({len(doc.paragraphs)} paragraf, {len(doc.tables)} tablo)")


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--md", type=Path, default=config.PROJECT_ROOT / "teslim" / "TEKNIK_RAPOR.md")
    p.add_argument("--out", type=Path, default=None)
    args = p.parse_args()
    convert(args.md, args.out or args.md.with_suffix(".docx"))


if __name__ == "__main__":
    main()
